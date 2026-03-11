"""
DOCX Resume Extraction Utility
------------------------------------------------
Extracts all visible text from DOCX files, including paragraphs and
table cells. Normalizes whitespace and bullet formatting for better
LLM processing.
"""

import re
from docx import Document


def extract_docx(path: str) -> str:
    doc = Document(path)
    parts = []

    # -----------------------------------------
    # Extract paragraphs
    # -----------------------------------------
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            text = _normalize_bullets(text)
            parts.append(text)

    # -----------------------------------------
    # Extract tables
    # -----------------------------------------
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    cell_text = _normalize_bullets(cell_text)
                    parts.append(cell_text)

    combined = "\n".join(parts)
    combined = _clean_docx_text(combined)
    return combined.strip()


def _normalize_bullets(text: str) -> str:
    """Convert DOCX bullets and unicode bullets to '- ' for LLM clarity."""
    bullet_chars = ["•", "◦", "▪", "–", "—", "·"]
    for b in bullet_chars:
        text = text.replace(b, "- ")
    text = re.sub(r"^\s*-\s*", "- ", text)
    return text


def _clean_docx_text(text: str) -> str:
    """Normalize spacing and remove DOCX artifacts."""
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"[ ]{2,}", " ", text)
    lines = [ln.strip() for ln in text.splitlines()]
    return "\n".join(lines)