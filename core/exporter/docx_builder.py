"""
Exports tailored resume text into DOCX with
resume-friendly formatting (compact spacing, headings, bullets).
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.style import WD_STYLE_TYPE


def export_to_docx(text: str, output_path: str):
    doc = Document()

    # Page setup (margins)
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    # Base normal style (body text)
    normal_style = doc.styles["Normal"]
    normal_font = normal_style.font
    normal_font.name = "Calibri"
    normal_font.size = Pt(11)

    normal_pf = normal_style.paragraph_format
    normal_pf.space_before = Pt(0)
    normal_pf.space_after = Pt(2)
    normal_pf.line_spacing = 1.0

    # Custom heading style for sections
    styles = doc.styles
    if "Resume Heading" in styles:
        heading_style = styles["Resume Heading"]
    else:
        heading_style = styles.add_style("Resume Heading", WD_STYLE_TYPE.PARAGRAPH)
        h_font = heading_style.font
        h_font.name = "Calibri"
        h_font.size = Pt(11)
        h_font.bold = True
        h_font.all_caps = True

        h_pf = heading_style.paragraph_format
        h_pf.space_before = Pt(6)
        h_pf.space_after = Pt(2)
        h_pf.line_spacing = 1.0

    # Bullet Style
    bullet_style_name = "List Bullet"
    if bullet_style_name in styles:
        bullet_style = styles[bullet_style_name]
        # Tighten the bullets
        b_pf = bullet_style.paragraph_format
        b_pf.space_before = Pt(0)
        b_pf.space_after = Pt(0)
        b_pf.line_spacing = 1.0
    else:
        bullet_style = normal_style  # Fallback

    # Helper Functions
    def is_section_heading(line: str) -> bool:
        # Heuristic: Short, mostly uppercase, not ending with a period.
        stripped = line.strip()
        if not stripped:
            return False

        if len(stripped) > 40:
            return False

        if stripped.endswith("."):
            return False

        # Allow spaces but require mostly uppercase letters
        letters = [ch for ch in stripped if ch.isalpha()]
        if not letters:
            return False

        upper_ratio = sum(ch.isupper() for ch in letters) / len(letters)
        return upper_ratio > 0.7

    def is_bullet(line: str) -> bool:
        stripped = line.lstrip()
        return stripped.startswith(("-", "•", "*"))

    # Normalize excessive blank lines from the model
    lines = text.splitlines()
    last_was_blank = False

    for raw_line in lines:
        line = raw_line.rstrip()

        # Handle blank lines (limit consecutive blanks)
        if not line.strip():
            if not last_was_blank:
                doc.add_paragraph("")  # Single empty line
            last_was_blank = True
            continue

        last_was_blank = False
        stripped = line.strip()

        # Section heading?
        if is_section_heading(stripped):
            p = doc.add_paragraph(stripped)
            p.style = heading_style
            continue

        # Bullet point?
        if is_bullet(stripped):
            # Remove leading bullet symbols and spaces
            for prefix in ("- ", "• ", "* "):
                if stripped.startswith(prefix):
                    stripped = stripped[len(prefix) :].strip()
                    break
            p = doc.add_paragraph(
                stripped,
                style=bullet_style_name if bullet_style_name in styles else None,
            )
            continue

        # Regular body text
        p = doc.add_paragraph(stripped)
        p.style = normal_style

    doc.save(output_path)
