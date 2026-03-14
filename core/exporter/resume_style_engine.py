"""
ResumeStyleEngine
------------------

Takes plain tailored resume text (from ResumeTailor.generate()) and
renders it as a formatted DOCX using the selected style template.

Supported styles:
    "swiss"         — Arial, minimal, no color, all-caps headings
    "spearmint"     — Calibri, spearmint green accent, left-border headings
    "coral"         — Georgia serif, coral/salmon accent, centered header
    "modern_writer" — Calibri, navy accent, spaced-caps name, full-width rule
    "prestige"       — Calibri, centered header, black rules, your personal style

Usage:
    from core.exporter.resume_style_engine import ResumeStyleEngine
    engine = ResumeStyleEngine()
    engine.export(tailored_text, style="spearmint", output_path="Resume.docx")
"""

import re
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── Style definitions ────────────────────────────────────────────────────────
STYLES = {
    "swiss": {
        "label":        "Swiss",
        "font":         "Arial",
        "name_size":    18,
        "body_size":    9.5,
        "heading_size": 9,
        "accent":       None,
        "heading_caps": True,
        "header_align": "left",
        "name_color":   "111111",
        "body_color":   "222222",
        "heading_color":"111111",
        "rule_color":   "AAAAAA",
        "company_color":"555555",
    },
    "spearmint": {
        "label":        "Spearmint",
        "font":         "Calibri",
        "name_size":    22,
        "body_size":    9.5,
        "heading_size": 11,
        "accent":       "2E8B6E",
        "heading_caps": True,
        "header_align": "left",
        "name_color":   "1A1A1A",
        "body_color":   "333333",
        "heading_color":"2E8B6E",
        "rule_color":   "2E8B6E",
        "company_color":"2E8B6E",
    },
    "coral": {
        "label":        "Coral",
        "font":         "Georgia",
        "name_size":    24,
        "body_size":    9.5,
        "heading_size": 11,
        "accent":       "C0533A",
        "heading_caps": False,
        "header_align": "center",
        "name_color":   "C0533A",
        "body_color":   "333333",
        "heading_color":"C0533A",
        "rule_color":   "C0533A",
        "company_color":"C0533A",
    },
    "modern_writer": {
        "label":        "Modern Writer",
        "font":         "Calibri",
        "name_size":    26,
        "body_size":    9.5,
        "heading_size": 10,
        "accent":       "1E3A5F",
        "heading_caps": True,
        "header_align": "left",
        "name_color":   "1E3A5F",
        "body_color":   "333333",
        "heading_color":"1E3A5F",
        "rule_color":   "1E3A5F",
        "company_color":"555555",
    },
    "prestige": {
        "label":        "Antonio",
        "font":         "Calibri",
        "name_size":    20,
        "body_size":    9.5,
        "heading_size": 11,
        "accent":       None,
        "heading_caps": True,
        "header_align": "center",
        "name_color":   "000000",
        "body_color":   "222222",
        "heading_color":"000000",
        "rule_color":   "888888",
        "company_color":"444444",
    },
}

DEFAULT_STYLE = "prestige"

# Common section header keywords for detection
SECTION_KEYWORDS = {
    "experience", "work experience", "professional experience",
    "education", "skills", "summary", "profile", "objective",
    "projects", "certifications", "achievements", "awards",
    "leadership", "technical skills", "core competencies",
    "volunteer", "recognitions", "awards / recognitions / volunteer work",
    "awards & recognition",
}


# ── XML helpers ──────────────────────────────────────────────────────────────
def _hex_to_rgb(hex_color: str):
    h = hex_color.lstrip("#")
    return tuple(int(h[i:i+2], 16) for i in (0, 2, 4))


def _set_paragraph_border_bottom(para, color="888888", size=6):
    """Add a bottom border to a paragraph via direct XML."""
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color.lstrip("#"))
    pBdr.append(bottom)
    pPr.append(pBdr)


def _set_paragraph_border_left(para, color="2E8B6E", size=18):
    """Add a left border bar to a paragraph (Spearmint headings)."""
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(size))
    left.set(qn("w:space"), "4")
    left.set(qn("w:color"), color.lstrip("#"))
    pBdr.append(left)
    pPr.append(pBdr)


def _add_tab_stop_right(para, position_dxa=8640):
    """Add a right-aligned tab stop to a paragraph."""
    pPr = para._p.get_or_add_pPr()
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:pos"), str(position_dxa))
    tabs.append(tab)
    pPr.append(tabs)


def _set_run_color(run, hex_color: str):
    r, g, b = _hex_to_rgb(hex_color)
    run.font.color.rgb = RGBColor(r, g, b)


def _set_spacing(para, before=0, after=60, line=None):
    pf = para.paragraph_format
    pf.space_before = Pt(before / 20) if before else Pt(0)
    pf.space_after  = Pt(after / 20)
    if line:
        from docx.shared import Pt as _Pt
        pf.line_spacing = _Pt(line / 20)


# ── Resume text parser ───────────────────────────────────────────────────────
def _is_section_heading(line: str) -> bool:
    stripped = line.strip()
    if not stripped or len(stripped) > 70:
        return False
    low = stripped.lower()
    if low in SECTION_KEYWORDS:
        return True
    if stripped.isupper() and len(stripped.split()) <= 6:
        return True
    letters = [c for c in stripped if c.isalpha()]
    if letters and sum(c.isupper() for c in letters) / len(letters) > 0.65:
        return True
    return False


def _is_bullet(line: str) -> bool:
    return line.lstrip().startswith(("-", "•", "*", "–", "—"))


def _clean_bullet(line: str) -> str:
    s = line.lstrip()
    for prefix in ("- ", "• ", "* ", "– ", "— ", "-", "•", "*", "–", "—"):
        if s.startswith(prefix):
            return s[len(prefix):].strip()
    return s.strip()


# ── Main engine ──────────────────────────────────────────────────────────────
class ResumeStyleEngine:

    def export(self, text: str, style: str = DEFAULT_STYLE,
               output_path: str = "Tailored_Resume.docx") -> str:
        """
        Render tailored resume text into a formatted DOCX.
        Returns the output path on success.
        """
        cfg = STYLES.get(style, STYLES[DEFAULT_STYLE])
        doc = Document()

        # Page setup — US Letter, 0.75" margins
        section = doc.sections[0]
        section.page_width  = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin    = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin   = Inches(0.9)
        section.right_margin  = Inches(0.9)

        # Clear default styles
        normal = doc.styles["Normal"]
        normal.font.name = cfg["font"]
        normal.font.size = Pt(cfg["body_size"])
        normal.paragraph_format.space_before = Pt(0)
        normal.paragraph_format.space_after  = Pt(3)

        lines = text.splitlines()
        self._render(doc, lines, cfg)

        doc.save(output_path)
        return output_path

    # ------------------------------------------------------------------
    def _render(self, doc, lines, cfg):
        font       = cfg["font"]
        body_size  = cfg["body_size"]
        body_color = cfg["body_color"]
        tab_pos    = 8500   # right tab position in DXA (~5.9 inches from left)

        # Detect header block (first non-empty lines before first section heading)
        header_lines = []
        body_start   = 0
        for i, line in enumerate(lines):
            if _is_section_heading(line):
                body_start = i
                break
            if line.strip():
                header_lines.append(line.strip())

        # ── Render header ────────────────────────────────────────────
        align = (WD_ALIGN_PARAGRAPH.CENTER
                 if cfg["header_align"] == "center"
                 else WD_ALIGN_PARAGRAPH.LEFT)

        for idx, hline in enumerate(header_lines):
            p = doc.add_paragraph()
            p.alignment = align
            _set_spacing(p, before=0, after=40 if idx < len(header_lines) - 1 else 80)

            if idx == 0:
                # Name line
                run = p.add_run(hline)
                run.font.name  = font
                run.font.size  = Pt(cfg["name_size"])
                run.font.bold  = True
                _set_run_color(run, cfg["name_color"])
            else:
                run = p.add_run(hline)
                run.font.name  = font
                run.font.size  = Pt(body_size + 0.5)
                run.font.bold  = False
                _set_run_color(run, "444444")

        # Rule under header
        rule_para = doc.add_paragraph()
        _set_paragraph_border_bottom(rule_para, cfg["rule_color"], size=8
                                     if cfg["name_color"] != "000000" else 12)
        _set_spacing(rule_para, before=0, after=120)

        # ── Render body ──────────────────────────────────────────────
        i = body_start
        while i < len(lines):
            line = lines[i]
            stripped = line.strip()

            if not stripped:
                i += 1
                continue

            # Section heading
            if _is_section_heading(stripped):
                p = doc.add_paragraph()
                heading_text = (stripped.upper()
                                if cfg["heading_caps"] else stripped.title())
                run = p.add_run(heading_text)
                run.font.name  = font
                run.font.size  = Pt(cfg["heading_size"])
                run.font.bold  = True
                _set_run_color(run, cfg["heading_color"])
                _set_spacing(p, before=12, after=6)

                # Style-specific heading decoration
                if cfg.get("accent") and cfg["header_align"] != "center":
                    # Spearmint: left border bar
                    if cfg["accent"] == "2E8B6E":
                        _set_paragraph_border_left(p, cfg["accent"])
                    else:
                        # Coral, Modern Writer: bottom rule in accent
                        _set_paragraph_border_bottom(p, cfg["rule_color"], size=8)
                else:
                    # Swiss, Antonio: grey bottom rule
                    _set_paragraph_border_bottom(p, cfg["rule_color"], size=6)

                i += 1
                continue

            # Bullet point
            if _is_bullet(stripped):
                p = doc.add_paragraph(style="List Bullet")
                run = p.add_run(_clean_bullet(stripped))
                run.font.name  = font
                run.font.size  = Pt(body_size)
                _set_run_color(run, body_color)
                _set_spacing(p, before=0, after=30)
                i += 1
                continue

            # Job/edu title line with right-aligned date (detect tab or date pattern)
            date_pattern = re.search(
                r"(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec|January|February|"
                r"March|April|June|July|August|September|October|November|December"
                r"|\d{4})[^|]*?(\d{4}|Present|Current)", stripped
            )
            next_line = lines[i + 1].strip() if i + 1 < len(lines) else ""
            is_job_title = (
                date_pattern and
                not _is_bullet(stripped) and
                len(stripped) < 100
            )

            if is_job_title:
                # Split title from date
                date_match = re.search(
                    r"(\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec|"
                    r"January|February|March|April|June|July|August|September|"
                    r"October|November|December|\d{4}).{0,40}"
                    r"(?:\d{4}|Present|Current)\b)", stripped
                )
                if date_match:
                    date_str   = date_match.group(0).strip()
                    title_str  = stripped[:date_match.start()].strip(" –—|-\t")
                else:
                    title_str = stripped
                    date_str  = ""

                p = doc.add_paragraph()
                _add_tab_stop_right(p, tab_pos)
                run_title = p.add_run(title_str)
                run_title.font.name  = font
                run_title.font.size  = Pt(body_size + 1)
                run_title.font.bold  = True
                _set_run_color(run_title, cfg.get("name_color", "000000"))

                if date_str:
                    run_tab  = p.add_run("\t")
                    run_date = p.add_run(date_str)
                    run_date.font.name    = font
                    run_date.font.size    = Pt(body_size)
                    run_date.font.italic  = True
                    _set_run_color(run_date, "666666")

                _set_spacing(p, before=8, after=20)

                # Check if next line is company/location (italic muted)
                if next_line and not _is_section_heading(next_line) and not _is_bullet(next_line):
                    i += 1
                    company_line = lines[i].strip()
                    cp = doc.add_paragraph()
                    run_co = cp.add_run(company_line)
                    run_co.font.name    = font
                    run_co.font.size    = Pt(body_size)
                    run_co.font.italic  = True
                    _set_run_color(run_co, cfg["company_color"])
                    _set_spacing(cp, before=0, after=30)

                i += 1
                continue

            # Regular text paragraph
            p = doc.add_paragraph()
            run = p.add_run(stripped)
            run.font.name  = font
            run.font.size  = Pt(body_size)
            _set_run_color(run, body_color)
            _set_spacing(p, before=0, after=40)
            i += 1