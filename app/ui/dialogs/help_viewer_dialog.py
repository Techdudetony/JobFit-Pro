"""
HelpViewerDialog
-----------------

In-app viewer for the JobFit Pro User Guide.

Layout:
  Left  — Table of contents (clickable section list)
  Right — QTextBrowser rendering the guide as styled HTML

The DOCX is parsed via python-docx and converted to HTML on first open,
then cached for the session. Falls back to a plain text render if parsing fails.

Accessed via Help → User Guide in the menu bar.
"""

import os
import re

from PyQt6.QtWidgets import (
    QDialog, QHBoxLayout, QVBoxLayout, QLabel,
    QPushButton, QTextBrowser, QListWidget,
    QListWidgetItem, QSplitter, QWidget, QFrame,
)
from PyQt6.QtCore import Qt, QUrl
from PyQt6.QtGui import QDesktopServices, QFont


# Path to the guide — relative to project root (works both dev and packaged)
def _guide_path() -> str:
    import sys
    # PyInstaller sets sys._MEIPASS for bundled resources
    base = getattr(sys, "_MEIPASS", None) or os.getcwd()
    return os.path.join(base, "assets", "docs", "JobFitPro_UserGuide.docx")


# ── DOCX → HTML converter ────────────────────────────────────────────────────
def _docx_to_html(path: str) -> tuple[str, list[tuple[str, str]]]:
    """
    Parse the DOCX and return (html_string, toc_entries).
    toc_entries is a list of (anchor_id, display_text) for the sidebar.
    """
    from docx import Document
    from docx.shared import RGBColor

    doc   = Document(path)
    parts = []
    toc   = []
    h1_count = 0

    ACCENT = "#22C55E"

    parts.append("""
    <html><head><style>
      body  { font-family: Calibri, sans-serif; font-size: 11pt;
              color: #1E293B; line-height: 1.6; padding: 24px 32px; }
      h1    { font-size: 18pt; font-weight: 700; color: #22C55E;
              border-bottom: 2px solid #22C55E; padding-bottom: 4px;
              margin-top: 32px; margin-bottom: 12px; }
      h2    { font-size: 13pt; font-weight: 700; color: #1E293B;
              margin-top: 24px; margin-bottom: 8px; }
      h3    { font-size: 11pt; font-weight: 700; color: #334155;
              margin-top: 18px; margin-bottom: 6px; }
      p     { margin: 4px 0 8px 0; }
      ul    { margin: 4px 0 8px 24px; padding: 0; }
      li    { margin-bottom: 4px; }
      .tip  { background: #F0FDF4; border-left: 3px solid #22C55E;
              padding: 8px 12px; margin: 8px 0; border-radius: 4px;
              font-style: italic; color: #166534; }
      .note { background: #FFFBEB; border-left: 3px solid #F59E0B;
              padding: 8px 12px; margin: 8px 0; border-radius: 4px;
              font-style: italic; color: #92400E; }
      .shortcut { font-family: monospace; font-weight: bold;
                  color: #22C55E; background: #F0FDF4;
                  padding: 1px 5px; border-radius: 3px; }
      hr    { border: none; border-top: 1px solid #E2E8F0; margin: 20px 0; }
    </style></head><body>
    """)

    in_list = False

    # Skip cover-page paragraphs — find the first real H1 and start there
    first_h1_idx = next(
        (i for i, p in enumerate(doc.paragraphs)
         if p.style and "Heading 1" in p.style.name),
        0
    )

    for para in doc.paragraphs[first_h1_idx:]:
        text    = para.text.strip()
        style   = para.style.name if para.style else ""
        if not text:
            if in_list:
                parts.append("</ul>")
                in_list = False
            parts.append("<p>&nbsp;</p>")
            continue

        # Headings
        if "Heading 1" in style:
            if in_list:
                parts.append("</ul>"); in_list = False
            h1_count += 1
            anchor = f"section_{h1_count}"
            toc.append((anchor, text))
            parts.append(f'<h1 id="{anchor}">{_escape(text)}</h1>')

        elif "Heading 2" in style:
            if in_list:
                parts.append("</ul>"); in_list = False
            parts.append(f"<h2>{_escape(text)}</h2>")

        elif "Heading 3" in style:
            if in_list:
                parts.append("</ul>"); in_list = False
            parts.append(f"<h3>{_escape(text)}</h3>")

        # Tip / Note callouts
        elif text.startswith("💡") or "Tip:" in text:
            if in_list:
                parts.append("</ul>"); in_list = False
            parts.append(f'<div class="tip">{_escape(text)}</div>')

        elif text.startswith("📝") or "Note:" in text:
            if in_list:
                parts.append("</ul>"); in_list = False
            parts.append(f'<div class="note">{_escape(text)}</div>')

        # Bullets / numbered lists
        elif "List" in style or _is_bullet(text):
            if not in_list:
                parts.append("<ul>"); in_list = True
            clean = _clean_bullet(text)
            # Detect shortcut lines (contain +)
            if re.match(r"^Ctrl\s*[+\+]", clean):
                parts.append(f'<li><span class="shortcut">{_escape(clean.split("—")[0].strip())}</span>'
                              f' — {_escape(clean.split("—")[-1].strip())}</li>')
            else:
                parts.append(f"<li>{_render_runs(para)}</li>")

        # Horizontal rule paragraph
        elif set(text) <= set("─—-_"):
            if in_list:
                parts.append("</ul>"); in_list = False
            parts.append("<hr>")

        else:
            if in_list:
                parts.append("</ul>"); in_list = False
            parts.append(f"<p>{_render_runs(para)}</p>")

    if in_list:
        parts.append("</ul>")
    parts.append("</body></html>")

    return "".join(parts), toc


def _escape(text: str) -> str:
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def _is_bullet(text: str) -> bool:
    return bool(re.match(r"^[\•\-\–\—\*]\s", text))


def _clean_bullet(text: str) -> str:
    return re.sub(r"^[\•\-\–\—\*]\s+", "", text).strip()


def _render_runs(para) -> str:
    """Convert paragraph runs to HTML with bold/italic."""
    parts = []
    for run in para.runs:
        t = _escape(run.text)
        if run.bold and run.italic:
            t = f"<b><i>{t}</i></b>"
        elif run.bold:
            t = f"<b>{t}</b>"
        elif run.italic:
            t = f"<i>{t}</i>"
        parts.append(t)
    return "".join(parts) or _escape(para.text)


# ── Dialog ────────────────────────────────────────────────────────────────────
class HelpViewerDialog(QDialog):
    _cached_html: str = ""
    _cached_toc:  list = []

    def __init__(self, parent=None):
        super().__init__(parent)
        self.setObjectName("HelpViewerDialog")
        self.setWindowTitle("JobFit Pro — User Guide")
        self.setMinimumSize(960, 680)
        self.setModal(False)   # non-modal so user can follow along
        self.setAttribute(Qt.WidgetAttribute.WA_DeleteOnClose, False)

        # Set window icon
        import os, sys
        from PyQt6.QtGui import QIcon
        base = getattr(sys, "_MEIPASS", None) or os.getcwd()
        icon_path = os.path.join(base, "assets", "desktop_icon.png")
        if os.path.exists(icon_path):
            self.setWindowIcon(QIcon(icon_path))
        self._build_ui()
        self._load_guide()

    # ── UI ────────────────────────────────────────────────────────────────────
    def _build_ui(self):
        root = QVBoxLayout(self)
        root.setContentsMargins(0, 0, 0, 0)
        root.setSpacing(0)

        # Header bar — compact, fixed height
        header = QWidget()
        header.setObjectName("helpHeader")
        header.setAttribute(Qt.WidgetAttribute.WA_StyledBackground, True)
        header.setFixedHeight(44)
        h_layout = QHBoxLayout(header)
        h_layout.setContentsMargins(16, 6, 16, 6)

        self.btn_external = QPushButton("↗ Open in Word")
        self.btn_external.setProperty("panelButton", True)
        self.btn_external.setFixedHeight(28)
        self.btn_external.clicked.connect(self._open_externally)

        btn_close = QPushButton("Close")
        btn_close.setFixedHeight(28)
        btn_close.clicked.connect(self.close)

        h_layout.addStretch()
        h_layout.addWidget(self.btn_external)
        h_layout.addWidget(btn_close)

        root.addWidget(header)

        # Divider
        div = QFrame()
        div.setFrameShape(QFrame.Shape.HLine)
        div.setObjectName("dialogDivider")
        root.addWidget(div)

        # Splitter: TOC left, content right
        splitter = QSplitter(Qt.Orientation.Horizontal)
        splitter.setObjectName("helpSplitter")

        # Table of contents
        toc_widget = QWidget()
        toc_widget.setObjectName("helpToc")
        toc_widget.setAttribute(Qt.WidgetAttribute.WA_StyledBackground, True)
        toc_layout = QVBoxLayout(toc_widget)
        toc_layout.setContentsMargins(0, 0, 0, 0)
        toc_layout.setSpacing(0)

        toc_header = QLabel("  Contents")
        toc_header.setObjectName("helpTocHeader")
        toc_header.setFixedHeight(36)
        toc_layout.addWidget(toc_header)

        self.toc_list = QListWidget()
        self.toc_list.setObjectName("helpTocList")
        self.toc_list.setFrameShape(QFrame.Shape.NoFrame)
        self.toc_list.itemClicked.connect(self._on_toc_clicked)
        toc_layout.addWidget(self.toc_list)

        # Content browser
        self.browser = QTextBrowser()
        self.browser.setObjectName("helpBrowser")
        self.browser.setOpenLinks(False)
        self.browser.setFrameShape(QFrame.Shape.NoFrame)

        splitter.addWidget(toc_widget)
        splitter.addWidget(self.browser)
        splitter.setSizes([200, 760])
        splitter.setHandleWidth(1)

        root.addWidget(splitter)

    # ── Load ──────────────────────────────────────────────────────────────────
    def _load_guide(self):
        path = _guide_path()

        if not os.path.exists(path):
            self.browser.setHtml(
                "<h2 style='color:#F87171;'>Guide not found</h2>"
                f"<p>Expected at:<br><code>{path}</code></p>"
                "<p>Ensure <code>assets/docs/JobFitPro_UserGuide.docx</code> "
                "exists in the project root.</p>"
            )
            return

        # Use cached version if already parsed this session
        # Cache is invalidated when the class is reloaded (app restart)
        if HelpViewerDialog._cached_html and os.path.exists(path):
            self._apply(HelpViewerDialog._cached_html, HelpViewerDialog._cached_toc)
            return

        try:
            html, toc = _docx_to_html(path)
            HelpViewerDialog._cached_html = html
            HelpViewerDialog._cached_toc  = toc
            self._apply(html, toc)
        except Exception as e:
            self.browser.setHtml(
                f"<h2 style='color:#F87171;'>Could not parse guide</h2>"
                f"<p>{e}</p>"
                f"<p>You can still open it externally using the button above.</p>"
            )

    def _apply(self, html: str, toc: list):
        self.browser.setHtml(html)
        self.toc_list.clear()
        for anchor, text in toc:
            item = QListWidgetItem(text)
            item.setData(Qt.ItemDataRole.UserRole, anchor)
            self.toc_list.addItem(item)

    # ── Navigation ────────────────────────────────────────────────────────────
    def _on_toc_clicked(self, item: QListWidgetItem):
        anchor = item.data(Qt.ItemDataRole.UserRole)
        self.browser.scrollToAnchor(anchor)

    def _open_externally(self):
        path = _guide_path()
        if os.path.exists(path):
            QDesktopServices.openUrl(QUrl.fromLocalFile(path))
        else:
            from PyQt6.QtWidgets import QMessageBox
            QMessageBox.warning(self, "File Not Found",
                                f"Could not find:\n{path}")